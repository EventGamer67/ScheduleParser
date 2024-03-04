import urllib
import requests
from pdf2docx import Converter

import supbase
from urllib.request import urlopen
from bs4 import BeautifulSoup
import datetime
from models import Group, Course, Teacher, Cabinet
from supbase import addGroup, addCourse, addTeacher, addCabinet, getParaNameAndTeacher, addFullZamenaGroup

SCHEDULE_URL = 'https://www.uksivt.ru/zameny'
#SCHEDULE_URL = 'http://127.0.0.1:3000/c:/Users/Danil/Desktop/Uksivt/sample.html'

BASEURL = 'https://www.uksivt.ru/'
from typing import List
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table


def parseParas(filename: str, date, sup,data):
    # cv = Converter(f'{filename}.pdf')
    # cv.convert('schedule' + '.docx', start=0, end=None)
    # cv.close()
    doc: DocumentType = Document('schedule.docx')
    groups = []
    for i in doc.paragraphs:
        text = i.text
        if (text.__contains__("Группа - ")):
            filter = text.split(' ')
            groups.append(filter[-1])

    for i in groups:
        if get_group_by_id(target_name=i, sup=sup, groups=data.GROUPS,data=data):
            pass

    tables = []
    for i in doc.tables:
        tables.append(i)

    tables = extract_all_tables_to_rows(tables)
    fin = []
    temp = []
    for row in tables:
        try:
            if (len(row[0]) > 1):
                for w in row:
                    temp.append(w)
            else:
                temp.append(row)
        except:
            pass
        fin.append(i)

    temp = clearDiscipline(temp)
    temp = clearSingleStrings(temp)

    divided = defineGroups(groups, temp)
    for gruppa in divided:
        paras = divided[gruppa]
        divided[gruppa] = removeDuplicates(paras)
        paras = divided[gruppa]
        divided[gruppa] = removeDoubleRows(paras)
        paras = divided[gruppa]
        divided[gruppa] = recoverTeachers(paras)
        ParasGroupToSoup(group=get_group_by_id(target_name=gruppa, groups=data.GROUPS, sup=sup,data=data), paras=divided[gruppa],sup=sup, startday=date,data=data)
    pass


def parseZamenas(filename: str, date, sup, data):
    all_rows = get_all_tables(filename)
    workRows = []
    for i in all_rows:
        if (not is_nested(i)):
            workRows.extend(i)
        else:
            workRows.append(i)
    workRows = clearNonDataRows(workRows)
    workRows = clear_empty_sublists(workRows)
    workRows = remove_headers(workRows)

    fullzamenagroups = []

    iteration = 0
    for i in workRows:
        iteration = iteration + 1
        if (i[0] == ''):
            i[0] = workRows[iteration - 2][0]

    for i in workRows:
        if(i.count(i[0]) == len(i)):
            fullzamenagroups.append(i[0].strip().replace(' ',""))

    cleaned = []
    for i in workRows:
        i.pop(2)

    for i in workRows:
        i.pop(2)

    for i in workRows:
        if (i[0] == i[1] and i[2] == i[3]):
            workRows.remove(i)

    editet = []
    for i in workRows:
        paras = i[1].split(',')
        row = i.copy()
        if (len(paras) >= 2):
            for para in paras:
                new = row.copy()
                new[1] = para
                editet.append(new)
        else:
            editet.append(i)

    workRows = editet

    for row in workRows:
        group = get_group_by_id(data.GROUPS, row[0], sup=sup,data=data)
        if group is not None:
            row[0] = group.id

    for row in workRows:
        course = get_course_by_id(data.COURSES, row[2], sup=sup,data=data)
        if course is not None:
            row[2] = course.id

    for row in workRows:
        teacher = get_teacher_by_id(data.TEACHERS, row[3], sup=sup,data=data)
        if teacher is not None:
            row[3] = teacher.id

    for row in workRows:
        cabinet = get_cabinet_by_id(data.CABINETS, row[4], sup=sup,data=data)
        if cabinet is not None:
            row[4] = cabinet.id

    for i in workRows:
        supbase.addZamena(sup=sup, group=i[0], number=i[1], course=i[2], teacher=i[3], cabinet=i[4], date=date)
        pass

    for i in fullzamenagroups:
        addFullZamenaGroup(sup=sup, group=get_group_by_id(target_name=i,data=data,groups=data.GROUPS,sup=sup).id,date=date)
    pass


def removeDoubleRows(table):
    alreadyExist = []
    for row in table:
        if (not row in alreadyExist):
            alreadyExist.append(row)
    return alreadyExist


def removeDuplicates(table):
    index = 0
    cleared = []
    for row in table:
        if ((table[index - 1])[0] == row[0]):
            continue
        else:
            cleared.append(row)
            index = index + 1
    return cleared


def ParasGroupToSoup(group, paras, startday, sup,data):
    print()
    print(group)
    print()
    date = startday
    for para in paras:
        number = para[0]
        ParaMonday: str = para[1]
        paraTuesday: str = para[3]
        paraWednesday: str = para[5]
        paraThursday: str = para[7]
        paraFriday: str = para[9]
        paraSaturday: str = para[11]
        days = [ParaMonday, paraTuesday, paraWednesday, paraThursday, paraFriday, paraSaturday]

        loopindex = 0
        for day in days:
            aww = getParaNameAndTeacher(day)
            if aww is not None:
                teacher = get_teacher_by_id(target_name=aww[0], teachers=data.TEACHERS, sup=sup,data=data)
                course = get_course_by_id(target_name=aww[1], courses=data.COURSES, sup=sup,data=data)
                cabinet = get_cabinet_by_id(target_name=para[2*(loopindex+1)], cabinets=data.CABINETS, sup=sup,data=data)
                if (teacher is not None and course is not None and cabinet is not None):
                    supbase.addPara(group=group.id, number=number, teacher=teacher.id, cabinet=cabinet.id,course=course.id, date=str(date + datetime.timedelta(days=loopindex)), sup=sup)
                    pass
            loopindex = loopindex + 1
            pass
    pass


def recoverTeachers(table):
    aww = 0
    for row in table:
        if (row[0] == ''):
            if (len(row[1].split(' ')) > 2):
                text = table[aww - 1][1]
                table[aww - 1][1] = text + " \n" + row[1]
            if (len(row[3].split(' ')) > 2):
                text = table[aww - 1][3]
                table[aww - 1][3] = text + " \n" + row[3]
            if (len(row[5].split(' ')) > 2):
                text = table[aww - 1][5]
                table[aww - 1][5] = text + " \n" + row[5]
            if (len(row[7].split(' ')) > 2):
                text = table[aww - 1][7]
                table[aww - 1][7] = text + " \n" + row[7]
            if (len(row[9].split(' ')) > 2):
                text = table[aww - 1][9]
                table[aww - 1][9] = text + " \n" + row[9]
            if (len(row[11].split(' ')) > 2):
                text = table[aww - 1][11]
                table[aww - 1][11] = text + " \n" + row[11]
            table.remove(row)
        aww += 1
    return table


def defineGroups(groups, table):
    groupIndex = -1
    groupParas = []
    group = 'x'
    divided = {}
    for row in table:
        if (row[0] == '№'):
            divided[group] = groupParas
            groupIndex = groupIndex + 1
            group = groups[groupIndex]
            groupParas = []
        else:
            groupParas.append(row)
        pass
    else:
        divided[group] = groupParas
        pass
    divided.pop('x')
    return divided


def clearSingleStrings(table):
    cleared = []
    for row in table:
        if (not isinstance(row, str)):
            cleared.append(row)
    return cleared


def clearDiscipline(table):
    for row in table:
        if (len(row) >= 12):
            if (row[0] == '' and row[1] == '' and row[2] == '' and row[3] == '' and row[4] == '' and row[5] == '' and
                    row[6] == ''):
                table.remove(row)
                continue
            if (row[1].__contains__("Дисциплина, вид занятия, преподаватель")):
                table.remove(row)
                continue
            pass
    return table


def get_cabinet_by_id(cabinets, target_name, sup,data) -> Cabinet:
    for cabinet in cabinets:
        if cabinet.name == target_name:
            return cabinet
        else:
            continue
    try:
        addCabinet(target_name, sup=sup,data=data)
        return get_cabinet_by_id(cabinets=data.CABINETS, target_name=target_name, sup=sup,data=data)
    except:
        return None
    return None


def get_teacher_from_short_name(teachers: List[Teacher], shortName: str):
    for i in teachers:
        fio: List[str] = i.name.split(' ')
        if len(fio) > 2 and fio[0].strip() != '' and fio[1].strip() != '' and fio[2].strip() != '':
            compare_result = f"{fio[0]} {fio[1][0]}.{fio[2][0]}."
            if compare_result.lower().strip() == shortName.lower().strip():
                return i
    return None


def get_teacher_by_id(teachers, target_name, sup,data) -> Teacher:
    for teacher in teachers:
        if teacher.name == target_name:
            return teacher
        else:
            search = get_teacher_from_short_name(teachers=teachers, shortName=target_name)
            if search is not None:
                return search
    try:
        addTeacher(target_name, sup=sup,data=data)
        return get_teacher_by_id(teachers=data.TEACHERS, target_name=target_name, sup=sup,data=data)
    except:
        return None
    return None


def get_group_by_id(groups, target_name, sup,data) -> Group:
    for group in groups:
        if group.name == target_name:
            return group
        else:
            continue
    try:
        addGroup(target_name, sup=sup,data=data)
        return get_group_by_id(groups=data.GROUPS, target_name=target_name, sup=sup,data=data)
    except:
        return None
    return None


def get_course_by_id(courses, target_name, sup,data) -> Course:
    for course in courses:
        if course.name == target_name:
            return course
        else:
            continue
    try:
        addCourse(target_name, sup=sup,data=data)
        return get_course_by_id(data.COURSES, target_name=target_name, sup=sup,data=data)
    except:
        return None
    return None


def extract_all_tables_to_rows(tables: List[Table]) -> List[List[str]]:
    rows = []
    for table in tables:
        for row in table.rows:
            data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell.tables:
                    nested_table_rows = extract_all_tables_to_rows(cell.tables)
                    data.extend(nested_table_rows)
                else:
                    data.append(cell_text)
            rows.append(data)
    return rows


def is_nested(row: List[str]) -> bool:
    return isinstance(row, List) and all(isinstance(item, str) for item in row)


def remove_duplicates(input_list):
    unique_list = []
    for sublist in input_list:
        if sublist not in unique_list:
            unique_list.append(sublist)
    return unique_list


def check_family(i):
    if (len(i[1].split(' ')) > 2) or (len(i[3].split(' ')) > 2) or (len(i[5].split(' ')) > 2) or (
            len(i[7].split(' ')) > 2) or (len(i[9].split(' ')) > 2) or (len(i[11].split(' ')) > 2):
        if (i[2] == '' and i[4] == '' and i[6] == '' and i[8] == '' and i[10] == '' and i[12] == ''):
            return True
    return False


def remove_headers(rows):
    cleared = []
    for i in rows:
        if (i[0] != "Группа"):
            cleared.append(i)
    return cleared


def get_all_tables(filename: str) -> List[List[str]]:
    doc: DocumentType = Document(filename)
    tables = doc.tables
    all_rows = extract_all_tables_to_rows(tables)
    return all_rows


def clear_empty_sublists(nested_list: List[List[str]]) -> List[List[str]]:
    return [sublist for sublist in nested_list if any(item != '' for item in sublist)]


def clearNonDataRows(rows: List[List[str]]) -> List[List[str]]:
    filtered_rows = [row for row in rows if len(row) >= 7]
    return filtered_rows


def downloadFile(link: str, filename: str):
    response = requests.get(link)
    if response.status_code == 200:
        with open(filename, 'wb') as file:
            file.write(response.content)
        print(f"File '{filename}' has been downloaded successfully.")
    else:
        print("Failed to download the file.")


def getAllLinks(soup: BeautifulSoup):
    pass


def getLastZamenaLink(soup: BeautifulSoup):
    days = getMonthAvalibleDays(soup=soup, monthIndex=0)
    return urllib.parse.urljoin(BASEURL, getDaylink(soup=soup, monthIndex=0, day=days[-1]))


def getLastZamenaDate(soup: BeautifulSoup):
    days = getMonthAvalibleDays(soup=soup, monthIndex=0)
    month, year = str(getMonthsList(soup=soup)[0]).split(' ')
    date = datetime.date(2024, convertMonthNameToIndex(month) + 1, days[-1])
    return date


def getDaylink(soup: BeautifulSoup, monthIndex: int, day: int):
    table = getMonthTable(soup=soup, monthIndex=monthIndex)
    for link in table.find_all('a'):
        text = link.get_text()
        if (link):
            if (text.isdigit() and int(text) == day):
                return link.get('href')


def convertMonthNameToIndex(name):
    months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь',
              'Декабрь']
    return months.index(name)


def getMonthTable(soup: BeautifulSoup, monthIndex: int):
    newtables = soup.find_all('table', {'class': 'MsoNormalTable'})
    oldtables = soup.find_all('table', {'class': 'calendar-month'})
    newtables.extend(oldtables)
    return newtables[monthIndex]


def getAllMonthTables(soup: BeautifulSoup):
    newtables = soup.find_all('table', {'class': 'MsoNormalTable'})
    print(newtables)
    oldtables = soup.find_all('table', {'class': 'calendar-month'})
    newtables.extend(oldtables)
    return newtables


def getAllTablesLinks(tables):
    links = []
    for table in tables:
        tags = table.find_all('a')
        for tag in tags:
            text = tag.get_text()
            if (tag):
                if text.isdigit():
                    links.append( urllib.parse.urljoin("https://www.uksivt.ru/zameny/",tag.get('href')))

    return links


def getMonthAvalibleDays(soup: BeautifulSoup, monthIndex: int):
    days = []
    table = getMonthTable(soup=soup, monthIndex=monthIndex)
    links = table.find_all('a')
    for link in links:
        if (link.get_text().isdigit()):
            if (link):
                days.append(int(link.get_text()))
    return days


def getMonthsList(soup : BeautifulSoup):
    paragraphs_with_class = soup.find_all("p", class_="MsoNormal")
    list = []
    for par in paragraphs_with_class:
        paragraphText = par.get_text(strip=True)
        if (paragraphText != '' and paragraphText is not None and not paragraphText.isdigit() and paragraphText not in ['Пн','Вт','Ср','Чт','Пт','Сб','Вс']):
            list.append(par.get_text(strip=True))
    index = 0
    for month in list:
        if len(month.split(' ')) == 1:
            list[index] = f"{month} {list[index+1].split(' ')[1]}"
        index = index+1
    return list


def getLatestScheduleMonth(soup: BeautifulSoup) -> str:
    table = soup.find('table', {'class': 'ui-datepicker-calendar'})
    paragraphs_with_class = soup.find_all("p", class_="MsoNormal")

    for par in paragraphs_with_class:
        if (par.get_text(strip=True) != '' and par.get_text(strip=True) is not None):
            return par.get_text(strip=True).split()[0]


def getLatestSchedleFile():
    html = urlopen(SCHEDULE_URL).read()
    soup = BeautifulSoup(html, 'html.parser')

    table = soup.find('table', {'class': 'calendar-month'})
    paragraphs_with_class = soup.find_all("p", class_="MsoNormal")

    for par in paragraphs_with_class:
        if (par.get_text(strip=True) != '' and par.get_text(strip=True) is not None):
            print(par.get_text(strip=True).split()[0])
            break

    date = ''
    urls = []
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        for cell in cells:
            link = cell.find('a')
            if link:
                href = link.get('href')
                date = cell.get_text(strip=True)
                if date.isdigit():
                    print(f"Date: {date}, Href: {href}")
                    file_url = urllib.parse.urljoin(BASEURL, href)
                    urls.append(file_url)

    return urls[-1]
