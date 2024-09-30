"""
Модуль содержит в себе парсер замен и вспомогательные методы для него
"""


import re
from io import BytesIO
from docx import Document
from docx.table import Table
from datetime import date
from src.code.tools.functions import get_remote_file_hash
from src.code.models.cabinet_model import Cabinet
from src.code.models.course_model import Course
from src.code.models.group_model import Group
from src.code.models.teacher_model import Teacher
from src.code.models.zamena_table_model import ZamTable
from src.code.models.data_model import Data
from src.code.network.supabase_worker import SupaBaseWorker



def parseZamenas(stream: BytesIO, date_: date, data_model: Data, link:str, supabase_client: SupaBaseWorker):
    all_rows, header = _get_all_tables(stream)
    practice_groups = _extract_practice_groups(header, data_model)
    workRows = _prepare_work_rows(all_rows)  # Подговить все строки к дальнейшей фильтрации (см. след. строку)
    workRows = _filter_and_clean_rows(workRows)

    # Работа с готовымии элементами
    workRows, fullzamenagroups, liquidation = handle_special_cases(workRows, data_model)
    update_empty_group_column(workRows)
    workRows = process_multiple_entries(workRows)

    map_entities_to_ids(workRows, data_model, supabase_client)

    prepare_and_send_supabase_entries(workRows, practice_groups, liquidation, fullzamenagroups,
                                      date_, link, data_model, supabase_client)


def _extract_practice_groups(header, data_model: Data):
    """Extract practice groups from the header."""
    practice_groups = []
    for i in header:
        practice_groups.extend(SupaBaseWorker.get_groups_from_string(i.text, data_model=data_model))
    return practice_groups

def _prepare_work_rows(all_rows):
    """Prepare the rows for processing, flattening nested rows."""
    workRows = []
    for i in all_rows:
        workRows.extend(i if not _is_nested(i) else [i])
    return workRows

def _filter_and_clean_rows(workRows: list[str]):
    """Filter rows based on specific conditions and clean up data."""
    workRows = [i for i in workRows if len(i) == 7]
    for i in workRows:
        if i[2] == '' and i[3] != '':
            i[3] = ''
        if i[1] == '' and i[2] == '' and i[3] == '' and i[5] == '':
            i[4] = ''
    workRows = clearNonDataRows(workRows)
    workRows = clear_empty_sublists(workRows)
    workRows = remove_headers(workRows)
    workRows = removeDemoExam(workRows)
    return workRows

def handle_special_cases(workRows: list[str], data_model: Data):
    """Handle specific cases such as liquidation and removing duplicate data."""
    iteration = 0
    liquidation = list()
    fullzamenagroups = list()

    #set group name in zamena group rows
    for i in workRows:
        iteration += 1
        if i[0] == '':
            i[0] = workRows[iteration - 2][0]

    #Find full zamena groups
    for i in workRows:
        if i[0] != '' and len(set(i)) == 1:
            if "ликвидация" not in i[0].strip().lower():
                fullzamenagroups.append(i[0].strip().replace(' ',''))
                workRows.remove(i)
            else:
                try:
                    sample = i[0].strip().lower()
                    for gr in data_model.GROUPS:
                        if (gr.name.lower().strip() in sample):
                            liquidation.append(gr.id)
                            print(f"Ликвидация {gr.name}")
                except Exception as err:
                    print(err)
                    continue
                workRows.remove(i)


    # for i in workRows[:]:
    #     if i[0] == i[1] and i[2] == i[3]:
    #         if "ликвидация" in i[0].strip().lower():
    #             try:
    #                 sample = i[0].strip().lower()
    #                 for gr in data_model.GROUPS:
    #                     if (gr.name.lower().strip() in sample):
    #                         liquidation.append(gr.id)
    #                         print(f"Ликвидация {gr.name}")
    #             except Exception as err:
    #                 print(err)
    #                 continue
    #             workRows.remove(i)
    #         workRows.remove(i)


    print(20*"*")
    for i in workRows:
        print(i)

    print(20 * "*")

    # for i in workRows[:]:
    #     i.pop(2)
    #     i.pop(2)
    #     if all(cell == '' for cell in i[1:]):
    #         workRows.remove(i)
    #     elif len(set(i)) == 1:
    #         fullzamenagroups.append(i[0].strip().replace(' ', ""))
    #         workRows.remove(i)

    return workRows, fullzamenagroups, liquidation


def update_empty_group_column(workRows: list[str]):
    """
    Метод удаляет лишние тире, пробелы, запятые, точки и приводит к lowercase
    """
    for i in workRows:
        i[0] = i[0].replace(' ', '').replace(',', '').replace('.', '').lower()
        i[0] = re.sub(r'-{2,}', '-', i[0])

def process_multiple_entries(workRows: list[str]):
    """
    Выполняет обработку над нескольками строк путём удаления лишних запятых, точек и т.п.
    """
    editet = []
    for i in workRows:
        try:
            text = i[1].replace('.', ',')
            print(text)
            if text[-1] == ',':
                text = text[:-1]
            if text[0] == ',':
                text = text[1:]
            paras = text.split(',')
            if len(paras) >= 2:
                for para in paras:
                    new = i.copy()
                    new[1] = para
                    editet.append(new)
            else:
                editet.append(i)
        except Exception as error:
            print(error)
            print(i)
            continue
    return editet

def map_entities_to_ids(workRows: list, data_model: Data, supabase_client: SupaBaseWorker):
    """
    Применяет функцию к Ids и записывает их в строки workRows
    """
    for row in workRows:
        group = get_group_by_id(data_model.GROUPS, row[0], data_model, supabase_client)
        if group:
            row[0] = group.id

        course = get_course_by_id(data_model.COURSES, row[2], data_model, supabase_client)
        if course:
            row[2] = course.id

        teacher = get_teacher_by_id(data_model.TEACHERS, row[3], data_model, supabase_client)
        if teacher:
            row[3] = teacher.id

        course = get_course_by_id(data_model.COURSES, row[4], data_model, supabase_client)
        if course:
            row[4] = course.id

        teacher = get_teacher_by_id(data_model.TEACHERS, row[5], data_model, supabase_client)
        if teacher:
            row[5] = teacher.id

        cabinet = get_cabinet_by_id(data_model.CABINETS, row[6], data_model, supabase_client)
        if cabinet:
            row[6] = cabinet.id

def prepare_and_send_supabase_entries(workRows, practice_groups: list, liquidation: list, fullzamenagroups: list,
                                      date_: date, link, data_model: Data, supabase_client: SupaBaseWorker):
    """
    Подготовить данные и отправить в БД
    """
    practice_supabase = [{"group": i.id, 'date': str(date_)} for i in practice_groups]
    zamenas_supabase = [
        {"group": i[0], 'number': int(i[1]), 'course': i[4], 'teacher': i[5], 'cabinet': i[6], 'date': str(date_)}
        for i in workRows
    ]
    full_zamenas_groups = [
        {"group": get_group_by_id(target_name=i, data_model=data_model, groups=data_model.GROUPS, supabase_client=supabase_client).id, 'date': str(date_)}
        for i in fullzamenagroups
    ]
    liquidations = [{"group": i, 'date': str(date)} for i in liquidation]

    supabase_client.addZamenas(zamenas=zamenas_supabase)
    if full_zamenas_groups:
        supabase_client.addFullZamenaGroups(groups=full_zamenas_groups)
    if practice_supabase:
        supabase_client.add_practices(practices=practice_supabase)
    if liquidations:
        supabase_client.addLiquidations(liquidations=liquidations)

    hash = get_remote_file_hash(link)
    supabase_client.addNewZamenaFileLink(link, date=date_, hash=hash)


def _get_all_tables(stream: BytesIO):
    docx = Document(stream)
    tables = docx.tables
    all_rows = _extract_all_tables_to_rows(tables)
    return all_rows, docx.paragraphs


def _extract_all_tables_to_rows(tables: list[Table]) -> list[list[str]]:
    rows = []
    for table in tables:
        for row in table.rows:
            data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell.tables:
                    nested_table_rows = _extract_all_tables_to_rows(cell.tables)
                    data.extend(nested_table_rows)
                else:
                    data.append(cell_text)
            rows.append(data)
    return rows

def _is_nested(row: list[str]) -> bool:
    return isinstance(row, list) and all(isinstance(item, str) for item in row)


def clearNonDataRows(rows: list[list[str]]) -> list[list[str]]:
    filtered_rows = [row for row in rows if len(row) >= 7]
    return filtered_rows


def clear_empty_sublists(nested_list: list[list[str]]) -> list[list[str]]:
    return [sublist for sublist in nested_list if any(item != '' for item in sublist)]


def remove_headers(rows):
    cleared = []
    for i in rows:
        if (i[0] != "Группа"):
            cleared.append(i)
    return cleared


def removeDemoExam(rows):
    cleared = []
    for i in rows:
        if not i[0].__contains__("экзамен"):
            cleared.append(i)
    return cleared


def find_entity_by_name(entities, target_name, name_key='name', normalize_func=None):
    target_name_normalized = str(target_name).lower()
    if normalize_func:
        target_name_normalized = normalize_func(target_name_normalized)

    for entity in entities:
        entity_name_normalized = getattr(entity, name_key).lower()
        if entity_name_normalized == target_name_normalized:
            return entity
    return None


def add_and_get_entity(entity_type, add_func, entities, target_name, data_model):
    entity = find_entity_by_name(entities, target_name)
    if entity:
        return entity
    try:
        #raise Exception(f"not found {target_name}")
        add_func(target_name, data_model=data_model)
        entities = getattr(data_model, entity_type)
        return find_entity_by_name(entities, target_name)
    except Exception as e:
        print(f"Error adding {entity_type}: {e}")
        return None


def get_group_by_id(groups, target_name, data_model: Data, supabase_client: SupaBaseWorker) -> Group:
    return add_and_get_entity(
        entity_type='GROUPS',
        add_func=supabase_client.addGroup,
        entities=groups,
        target_name=target_name.replace('_', '-').lower(),
        data_model=data_model
    )


def get_cabinet_by_id(cabinets : list[Cabinet], target_name :str, data_model: Data, supabase_client: SupaBaseWorker) -> Cabinet:
    return add_and_get_entity(
        entity_type='CABINETS',
        add_func=supabase_client.addCabinet,
        entities=cabinets,
        target_name=target_name,
        data_model=data_model
    )


def get_course_by_id(courses, target_name, data_model: Data, supabase_client: SupaBaseWorker) -> Course:
    return add_and_get_entity(
        entity_type='COURSES',
        add_func=supabase_client.addCourse,
        entities=courses,
        target_name=target_name,
        data_model=data_model
    )


def get_teacher_by_id(teachers, target_name, data_model: Data, supabase_client: SupaBaseWorker) -> Teacher:
    teacher = find_entity_by_name(teachers, target_name)
    if teacher:
        return teacher

    teacher = get_teacher_from_short_name(teachers=teachers, shortName=target_name)
    if teacher:
        return teacher

    return add_and_get_entity(
        entity_type='TEACHERS',
        add_func=supabase_client.addTeacher,
        entities=teachers,
        target_name=target_name,
        data_model=data_model
    )


def get_teacher_from_short_name(teachers: list[Teacher], shortName: str):
    shortcomparer = shortName.replace('.', '').replace(',', '').replace(' ', '').lower().strip()

    for teacher in teachers:
        fio = teacher.name.split()
        if len(fio) < 3:
            continue

        # Формирование короткой версии имени преподавателя
        compare_result = f"{fio[0]}{fio[1][0]}{fio[2][0]}".lower().strip()

        if compare_result == shortcomparer or count_different_characters(compare_result, shortcomparer) == 1:
            print(f"taker {compare_result} and {shortcomparer}")
            print(f"set {teacher}")
            return teacher

    return None

def count_different_characters(str1, str2):
    if len(str1) != len(str2):
        return -1

    # Подсчет количества различных символов
    return sum(char1 != char2 for char1, char2 in zip(str1, str2))
